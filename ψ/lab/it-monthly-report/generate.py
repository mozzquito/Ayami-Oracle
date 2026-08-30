#!/usr/bin/env python3
"""Render an IT Service Monthly Report from a schema.yaml-shaped data file.

Usage:
    ./.venv/bin/python generate.py data/2026-08.yaml               # -> data/2026-08.md
    ./.venv/bin/python generate.py data/2026-08.yaml --docx         # also data/2026-08.docx

See README.md for the data-entry workflow (manual YAML vs. AI-assisted fill-in).
"""
import argparse
import copy
import re
import sys
from datetime import date, datetime
from pathlib import Path

import yaml
from jinja2 import ChainableUndefined, Environment, FileSystemLoader

HERE = Path(__file__).parent

STATUS_ICON = {
    "normal": "🟢 Normal", "ok": "🟢 OK", "green": "🟢",
    "warning": "🟡 Warning", "caution": "🟡 Caution", "yellow": "🟡",
    "critical": "🔴 Critical", "issue": "🔴 Issue", "red": "🔴",
}


def status_icon(value):
    if not value:
        return "⚠️ MISSING"
    return STATUS_ICON.get(str(value).lower(), str(value))


def delta(this_month, last_month):
    if this_month is None or last_month is None:
        return "-"
    d = round(this_month - last_month, 2)
    if d > 0:
        return f"▲ +{d}"
    if d < 0:
        return f"▼ {d}"
    return "= 0"


def _parse_date(value):
    if not value:
        return None
    if isinstance(value, (date, datetime)):
        return value if isinstance(value, date) else value.date()
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").date()
    except ValueError:
        return None


def license_status(expire, today=None):
    today = today or date.today()
    d = _parse_date(expire)
    if d is None:
        return "-" if not expire else "⚠️ invalid date"
    days = (d - today).days
    if days < 0:
        return f"🔴 EXPIRED ({-days}d ago)"
    if days <= 30:
        return f"🟡 expires in {days}d"
    return "🟢 OK"


def compute_license_alerts(data, today=None):
    today = today or date.today()
    alerts = list(data.get("software", {}).get("license_alerts") or [])
    for bucket in ("os_licenses", "office_licenses"):
        for item in data.get("software", {}).get(bucket) or []:
            name, expire = item.get("name"), item.get("expire")
            if not name:
                continue
            d = _parse_date(expire)
            if d is None:
                continue
            days = (d - today).days
            if days < 0:
                alerts.append(f"{name} หมดอายุไปแล้ว {-days} วัน ({expire})")
            elif days <= 30:
                alerts.append(f"{name} จะหมดอายุใน {days} วัน ({expire})")
    return alerts


PRIORITY_RANK = {"high": 3, "medium": 2, "low": 1}


def sort_recommendations(items):
    return sorted(items or [], key=lambda r: PRIORITY_RANK.get(str(r.get("priority", "")).lower(), 0),
                  reverse=True)


def find_previous_month(path: Path) -> Path | None:
    """Sibling YYYY-MM.yaml file that sorts immediately before this one.

    Only month files match — profile.yaml / history.yaml etc. are ignored so
    they can never be mistaken for the previous month's report.
    """
    month_file = re.compile(r"^\d{4}-\d{2}\.yaml$")
    siblings = sorted(
        p for p in path.parent.glob("*.yaml")
        if month_file.match(p.name) and p != path
    )
    earlier = [p for p in siblings if p.name < path.name]
    return earlier[-1] if earlier else None


def build_changelog(current: dict, prev: dict) -> list[str]:
    changes = []

    cur_count = (current.get("computers") or {}).get("contract_count")
    prev_count = (prev.get("computers") or {}).get("contract_count")
    if cur_count is not None and prev_count is not None and cur_count != prev_count:
        d = cur_count - prev_count
        changes.append(f"จำนวนคอมพิวเตอร์เปลี่ยนจาก {prev_count} เป็น {cur_count} ({'+' if d > 0 else ''}{d})")

    cur_titles = {a.get("title") for a in (current.get("health") or {}).get("critical_alerts") or [] if a.get("title")}
    prev_titles = {a.get("title") for a in (prev.get("health") or {}).get("critical_alerts") or [] if a.get("title")}
    for t in cur_titles - prev_titles:
        changes.append(f'Critical alert ใหม่: "{t}"')
    for t in prev_titles - cur_titles:
        changes.append(f'Critical alert หายไป: "{t}"')

    cur_up = (current.get("sla") or {}).get("uptime_percent")
    prev_up = (prev.get("sla") or {}).get("uptime_percent")
    if cur_up is not None and prev_up is not None and cur_up != prev_up:
        changes.append(f"SLA uptime เปลี่ยนจาก {prev_up}% เป็น {cur_up}%")

    cur_tickets = (current.get("tickets") or {})
    prev_tickets = (prev.get("tickets") or {})
    cur_total = (cur_tickets.get("incident_count") or 0) + (cur_tickets.get("service_request_count") or 0)
    prev_total = (prev_tickets.get("incident_count") or 0) + (prev_tickets.get("service_request_count") or 0)
    if cur_total != prev_total:
        changes.append(f"จำนวน ticket เปลี่ยนจาก {prev_total} เป็น {cur_total}")

    today = date.today()
    for bucket, label in (("os_licenses", "OS License"), ("office_licenses", "Office License")):
        prev_named = {i.get("name") for i in (prev.get("software") or {}).get(bucket) or [] if i.get("expire") and i.get("name")}
        for item in (current.get("software") or {}).get(bucket) or []:
            name, expire = item.get("name"), item.get("expire")
            d = _parse_date(expire)
            if not name or d is None or name in prev_named:
                continue
            days = (d - today).days
            if -30 <= days <= 30:
                changes.append(f'{label} "{name}" {"หมดอายุแล้ว" if days < 0 else f"จะหมดอายุใน {days} วัน"}')

    return changes


# Full skeleton of a report data file. merge_defaults() fills any missing
# key (without overriding existing values, explicit nulls included) so a
# partial/hand-edited YAML can never crash either renderer — QA-found bug:
# dropping e.g. `prepared_by` used to raise UndefinedError before the
# template's `or "MISSING"` fallbacks could help.
REPORT_DEFAULTS = {
    "report": {"title": "IT Service Monthly Report", "month": None, "year": None,
               "client_name": None, "report_date": None, "onsite_ma_date": None,
               "monitor_date": None},
    "prepared_by": {"name": None, "role": None, "phone": None, "email": None},
    "checked_by": {"name": None, "role": None, "phone": None, "email": None},
    "health": {"overall_status": None, "highlights": [], "critical_alerts": []},
    "sla": {"uptime_percent": None, "avg_response_hours": None,
            "avg_resolution_hours": None, "sla_met_percent": None},
    "trend_mom": {"ticket_count": {"this_month": None, "last_month": None},
                  "uptime_percent": {"this_month": None, "last_month": None},
                  "recurring_issues": []},
    "computers": {"contract_count": None, "serviced_count": None,
                  "by_type": {"desktop": 0, "laptop": 0, "all_in_one": 0, "macbook": 0},
                  "hard_disk": {"normal": 0, "caution": 0},
                  "battery": {"ok": 0, "degraded": 0},
                  "age_distribution": {"under_1y": 0, "1_2y": 0, "2_4y": 0, "5_7y": 0, "over_7y": 0},
                  "replacement_recommendation": None},
    "software": {"os_licenses": [], "office_licenses": [], "license_alerts": []},
    "server": {"contract_count": None, "serviced_count": None,
               "physical_servers": [], "guest_vms": [],
               "backup": {"scheduled_status": None, "last_restore_test": None}},
    "firewall_gateway": {"devices": [], "warning_critical_log": None},
    "network": {"devices": [], "warning_critical_log": None},
    "tickets": {"incident_count": 0, "service_request_count": 0,
                "status": {"in_progress": 0, "pending": 0, "done": 0}, "list": []},
    "scope_of_work": [],
    "recommendations": [],
    "sign_off": {"client_signer_name": None, "client_signed_date": None},
}


def merge_defaults(data: dict, defaults: dict) -> dict:
    """Recursively fill missing keys / null sections from defaults."""
    for key, val in defaults.items():
        if key not in data or data[key] is None:
            data[key] = copy.deepcopy(val)
        elif isinstance(val, dict) and isinstance(data[key], dict):
            merge_defaults(data[key], val)
    return data


def load_data(path: Path) -> dict:
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    data = merge_defaults(data, REPORT_DEFAULTS)
    data.setdefault("software", {})["license_alerts"] = compute_license_alerts(data)
    data["recommendations"] = sort_recommendations(data.get("recommendations"))

    trend = data.setdefault("trend_mom", {})
    tc = trend.setdefault("ticket_count", {})
    if tc.get("this_month") is None:
        cur_t = data.get("tickets") or {}
        tc["this_month"] = (cur_t.get("incident_count") or 0) + (cur_t.get("service_request_count") or 0)
    up = trend.setdefault("uptime_percent", {})
    if up.get("this_month") is None:
        up["this_month"] = (data.get("sla") or {}).get("uptime_percent")

    prev_path = find_previous_month(path)
    if prev_path is not None:
        prev_data = yaml.safe_load(prev_path.read_text(encoding="utf-8")) or {}
        if tc.get("last_month") is None:
            prev_t = prev_data.get("tickets") or {}
            tc["last_month"] = (prev_t.get("incident_count") or 0) + (prev_t.get("service_request_count") or 0)
        if up.get("last_month") is None:
            up["last_month"] = (prev_data.get("sla") or {}).get("uptime_percent")
        data["_changelog"] = build_changelog(data, prev_data)
        data["_changelog_prev_month"] = prev_path.stem
    else:
        data["_changelog"] = []
        data["_changelog_prev_month"] = None

    return data


def render_markdown(data: dict) -> str:
    env = Environment(
        loader=FileSystemLoader(str(HERE)),
        trim_blocks=True,
        lstrip_blocks=True,
        undefined=ChainableUndefined,
    )
    env.globals["_delta"] = delta
    env.globals["_license_status"] = license_status
    env.globals["_status_icon"] = status_icon
    template = env.get_template("template.md.j2")
    return template.render(**data)


def render_docx(data: dict, out_path: Path):
    """Wayama-style DOCX render — visual template cloned from the original
    Kittisampan July 2026 .docx: cover page, dotted index, two-tone section
    banners (dark blue / amber), Cordia New Thai body + Microsoft Sans Serif
    headings, logo header and a "Page X of Y" footer. Index page numbers
    assume one page per section (same convention as the original); a section
    that overflows shifts later numbers by one."""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor

    dark = RGBColor(0x32, 0x3E, 0x4F)   # section banners / headings
    amber = RGBColor(0xFF, 0xC0, 0x00)  # accent
    red = RGBColor(0xC0, 0x00, 0x00)
    ink = RGBColor(0x00, 0x20, 0x60)    # Thai body text
    gray = RGBColor(0xA6, 0xA6, 0xA6)
    en_font, th_font = "Microsoft Sans Serif", "Cordia New"
    width = Inches(6.27)                # A4 minus 1" side margins
    logo = HERE / "assets" / "wayama-logo.png"

    report = data.get("report", {})
    prepared = data.get("prepared_by", {})
    checked = data.get("checked_by", {})

    def style(run, font=th_font, size=14, color=ink, bold=False, italic=False):
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:cs"), font)           # complex-script font (Thai)
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(size * 2)))
        if bold:
            rPr.append(OxmlElement("w:bCs"))   # complex-script bold
        return run

    def fmt_date(v):
        """ISO -> '25 July 2026', matching the original report's display."""
        d = _parse_date(v)
        if d is not None:
            return f"{d.day} {d.strftime('%B %Y')}"
        return v or "-"

    def para(container, before=0, after=6, align=None):
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        return p

    def text(p, s, **kw):
        return style(p.add_run("" if s is None else str(s)), **kw)

    def field(p, instr, cached=None, **kw):
        r = style(p.add_run(), **kw)
        b = OxmlElement("w:fldChar")
        b.set(qn("w:fldCharType"), "begin")
        i = OxmlElement("w:instrText")
        i.set(qn("xml:space"), "preserve")
        i.text = f" {instr} "
        r._r.append(b)
        r._r.append(i)
        if cached is not None:
            sep = OxmlElement("w:fldChar")
            sep.set(qn("w:fldCharType"), "separate")
            t = OxmlElement("w:t")
            t.text = str(cached)
            r._r.append(sep)
            r._r.append(t)
        e = OxmlElement("w:fldChar")
        e.set(qn("w:fldCharType"), "end")
        r._r.append(e)

    def bottom_border(p, color="323E4F", sz="12"):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), color)
        pBdr.append(b)
        pPr.append(pBdr)

    def banner(doc, left, right, bookmark=None, bid=None):
        p = para(doc, before=4, after=14)
        p.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)
        text(p, left, font=en_font, size=22, color=dark, bold=True)
        text(p, "\t" + right, font=en_font, size=22, color=amber, bold=True)
        bottom_border(p)
        if bookmark:
            # Anchor for the Index page's PAGEREF fields -> real page numbers
            bs = OxmlElement("w:bookmarkStart")
            bs.set(qn("w:id"), str(bid))
            bs.set(qn("w:name"), bookmark)
            be = OxmlElement("w:bookmarkEnd")
            be.set(qn("w:id"), str(bid))
            p._p.get_or_add_pPr().addnext(bs)
            p._p.append(be)

    def sub_banner(doc, s):
        p = para(doc, before=10, after=6)
        text(p, s, font=th_font, size=15, color=dark, bold=True)

    def bullet(doc, s, **kw):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        text(p, s, **kw)

    def set_cell(cell, value, font=th_font, size=14, color=ink, bold=False,
                 align=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if align is not None:
            p.alignment = align
        text(p, "-" if value is None else value, font=font, size=size,
             color=color, bold=bold)

    def make_table(doc, headers, rows, widths, center_cols=(), size=14):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.autofit = False
        for i, h in enumerate(headers):
            set_cell(t.rows[0].cells[i], h, size=size, color=dark, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                set_cell(cells[i], v, size=size,
                         align=WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else None)
        for row in t.rows:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for i, w in enumerate(widths):
                row.cells[i].width = w
        t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        para(doc, after=4)
        return t

    def kv_table(doc, pairs, bordered=False):
        t = doc.add_table(rows=0, cols=2)
        t.autofit = False
        if bordered:
            t.style = "Table Grid"
        for k, v in pairs:
            r = t.add_row()
            r._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            cells = r.cells
            set_cell(cells[0], k, color=dark, bold=True)
            set_cell(cells[1], v)
            cells[0].width = Inches(2.4)
            cells[1].width = Inches(3.87)
        para(doc, after=4)
        return t

    # ---- Cover (section 1: side-band design — Wayama original, elevated) ----
    doc = Document()
    sec1 = doc.sections[0]
    sec1.page_width, sec1.page_height = Mm(210), Mm(297)
    sec1.left_margin = sec1.right_margin = Inches(1)
    sec1.top_margin = sec1.bottom_margin = Inches(0.8)

    band_h = Inches(9.9)  # exact row height; slack below for the trailing para
    cover = doc.add_table(rows=1, cols=2)
    cover.autofit = False
    row = cover.rows[0]
    row.height = band_h
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    band_cell, main_cell = row.cells
    band_cell.width = Inches(2.2)
    main_cell.width = Inches(4.07)

    def cell_margins(cell, top=0, left=0, bottom=0, right=0):
        """Set cell padding in twips (1/1440 inch)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("left", left),
                          ("bottom", bottom), ("right", right)):
            e = OxmlElement(f"w:{side}")
            e.set(qn("w:w"), str(val))
            e.set(qn("w:type"), "dxa")
            tcMar.append(e)
        tcPr.append(tcMar)

    def shade_cell(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    # Left band: datacenter photo edge-to-edge, thin amber separator on its
    # right edge; navy fallback if the photo is missing
    cell_margins(band_cell)
    tcPr = band_cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    rb = OxmlElement("w:right")
    rb.set(qn("w:val"), "single")
    rb.set(qn("w:sz"), "24")
    rb.set(qn("w:color"), "FFC000")
    tcB.append(rb)
    tcPr.append(tcB)
    photo = HERE / "assets" / "cover-datacenter.png"
    band_cell.text = ""
    bp = band_cell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(0)
    if photo.exists():
        pic_run = bp.add_run()
        pic_run.add_picture(str(photo), width=Inches(2.2), height=band_h)
        # crop the photo's sides (not distort) to the band's aspect ratio:
        # band 2.2x9.9 -> keep ~62% of the 369x1024 image width, centered
        blip = pic_run._element.xpath(".//a:blip")[0]
        src_rect = OxmlElement("a:srcRect")
        src_rect.set("l", "19170")
        src_rect.set("r", "19170")
        blip.addnext(src_rect)
    else:
        shade_cell(band_cell, "323E4F")

    # Right column: title / month / company / contacts / logo
    cell_margins(main_cell, top=0, left=504, bottom=0, right=100)
    main_cell.text = ""

    def cpara(before=0, after=4, align=None):
        p = main_cell.paragraphs[0] if not main_cell.paragraphs[0].runs and \
            not getattr(cpara, "_used", False) else main_cell.add_paragraph()
        cpara._used = True
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        return p

    title_words = (report.get("title") or "IT Service Monthly Report").split()
    line1 = " ".join(w.upper() for w in title_words[:len(title_words) // 2 or 1])
    line2 = " ".join(w.upper() for w in title_words[len(title_words) // 2 or 1:])

    p = cpara(before=64, after=0)
    text(p, line1, font=en_font, size=26, color=dark, bold=True)
    p = cpara(after=8)
    text(p, line2, font=en_font, size=26, color=dark, bold=True)
    p = cpara(after=14)
    text(p, f"{report.get('month', '')} {report.get('year', '')}",
         font=en_font, size=20, color=amber, bold=True)
    p = cpara(after=18)
    bottom_border(p, color="323E4F", sz="12")

    p = cpara(after=10)
    text(p, (report.get("client_name") or "⚠️ MISSING CLIENT_NAME").upper(),
         font=en_font, size=13, color=dark, bold=True)

    for label, person in (("PREPARED BY", prepared), ("CHECKED BY", checked),
                          ("CONTACT", prepared)):
        p = cpara(before=20, after=2)
        text(p, label, font=en_font, size=10, color=amber, bold=True)
        p = cpara(after=1)
        text(p, person.get("name") or "-", font=en_font, size=11,
             color=ink, bold=True)
        for ln in (person.get("role"), person.get("phone"), person.get("email")):
            p = cpara(after=1)
            text(p, ln or "-", font=en_font, size=10, color=ink)

    if logo.exists():
        p = cpara(before=48, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
        p.add_run().add_picture(str(logo), height=Inches(0.45))

    # Word requires a paragraph after a table; keep it 2pt so it can't spill
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(0)
    text(tp, " ", font=en_font, size=2)

    # ---- Content (section 2: logo header + client footer, continuous page numbers) ----
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    hdr = sec2.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(2)
    if logo.exists():
        hp.add_run().add_picture(str(logo), height=Inches(0.45))
    bottom_border(hp, color="FFC000", sz="8")

    ftr = sec2.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)
    text(fp, (report.get("client_name") or "").upper(), font=en_font, size=9,
         color=gray, bold=True)
    text(fp, "\tPage ", font=en_font, size=9, color=gray)
    field(fp, "PAGE", font=en_font, size=9, color=gray)
    text(fp, " of ", font=en_font, size=9, color=gray)
    field(fp, "NUMPAGES", font=en_font, size=9, color=gray)

    # ---- Index ----
    p = para(doc, before=4, after=14)
    p.paragraph_format.tab_stops.add_tab_stop(
        width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    text(p, "IT Service Monthly Report", font=en_font, size=22, color=dark, bold=True)
    text(p, "\tIndex", font=en_font, size=22, color=amber, bold=True)
    bottom_border(p)
    for i, (name, num, bm) in enumerate([
            ("Monitoring Report Summary", "03", "sec_summary"),
            ("Computer", "04", "sec_computer"),
            ("Software", "05", "sec_software"),
            ("Server", "06", "sec_server"),
            ("Firewall, Gateway", "07", "sec_firewall"),
            ("Network", "08", "sec_network"),
            ("Ticket Report Summary", "09", "sec_ticket"),
            ("Scope of Work Check", "10", "sec_scope"),
            ("Recommendations", "11", "sec_recs")]):
        c = dark if i % 2 == 0 else amber
        p = para(doc, before=6, after=6)
        p.paragraph_format.tab_stops.add_tab_stop(
            width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        text(p, name, font=en_font, size=13, color=c, bold=True)
        text(p, "\t", font=en_font, size=13, color=c, bold=True)
        # cached number shows immediately; PAGEREF re-resolves to the real
        # page when Word updates fields (print / select-all+F9), so overflow
        # from added data can't leave stale numbers behind
        field(p, f"PAGEREF {bm} \\h", cached=num,
              font=en_font, size=13, color=c, bold=True)
    doc.add_page_break()

    # ---- 1. Summary ----
    banner(doc, "Monitoring Report", "Summary", bookmark="sec_summary", bid=101)
    kv_table(doc, [
        ("Report Date (IT Maintenance)", fmt_date(report.get("report_date"))),
        ("Onsite MA", fmt_date(report.get("onsite_ma_date"))),
        ("Server Network Monitor", fmt_date(report.get("monitor_date"))),
    ], bordered=True)

    health = data.get("health", {})
    status_map = {
        "green": ("🟢 ปกติ (Normal)", ink),
        "yellow": ("🟡 ต้องเฝ้าระวัง (Warning)", dark),
        "red": ("🔴 วิกฤต (Critical)", red),
    }
    label, color = status_map.get(health.get("overall_status"),
                                  ("⚠️ NOT SET", red))
    p = para(doc, before=8, after=8)
    text(p, "Overall status: ", font=en_font, size=14, color=dark, bold=True)
    text(p, label, size=14, color=color, bold=True)

    sub_banner(doc, "Highlights")
    for h in health.get("highlights") or []:
        if h:
            bullet(doc, h)

    alerts = [a for a in health.get("critical_alerts") or [] if a.get("title")]
    if alerts:
        sub_banner(doc, "⚠️ Critical Alerts — ต้องการการตัดสินใจจากลูกค้า")
        make_table(doc, ["Severity", "Issue", "Detail", "Recommendation"],
                   [[("🔴" if a.get("severity") == "red" else "🟡"),
                     a.get("title"), a.get("description") or "-",
                     a.get("recommendation") or "⚠️ MISSING"] for a in alerts],
                   widths=[Inches(0.7), Inches(1.45), Inches(2.25), Inches(1.87)],
                   center_cols=(0,))

    sla = data.get("sla", {})
    sub_banner(doc, "SLA Performance")
    make_table(doc, ["Metric", "Value"], [
        ["Uptime %", sla.get("uptime_percent")],
        ["Avg. Response Time (hrs)", sla.get("avg_response_hours")],
        ["Avg. Resolution Time (hrs)", sla.get("avg_resolution_hours")],
        ["SLA Met %", sla.get("sla_met_percent")],
    ], widths=[Inches(3.5), Inches(2.77)], center_cols=(1,))

    trend = data.get("trend_mom", {})
    tc = trend.get("ticket_count") or {}
    up = trend.get("uptime_percent") or {}
    sub_banner(doc, "Trend (Month-over-Month)")
    make_table(doc, ["Metric", "This Month", "Last Month", "Δ"], [
        ["Ticket count", tc.get("this_month"), tc.get("last_month"),
         delta(tc.get("this_month"), tc.get("last_month"))],
        ["Uptime %", up.get("this_month"), up.get("last_month"),
         delta(up.get("this_month"), up.get("last_month"))],
    ], widths=[Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.27)],
        center_cols=(1, 2, 3))
    if trend.get("recurring_issues"):
        p = para(doc)
        text(p, "Recurring issues: ", size=14, color=dark, bold=True)
        text(p, ", ".join(trend["recurring_issues"]), size=14)
    if data.get("_changelog"):
        sub_banner(doc, f"Changes since {data.get('_changelog_prev_month')}")
        for c in data["_changelog"]:
            bullet(doc, c)
    doc.add_page_break()

    # ---- 2. Computer (original "Maintenance Computer" table shape) ----
    banner(doc, "Monitoring Report", "Computer", bookmark="sec_computer", bid=102)
    sub_banner(doc, "Maintenance Computer")
    c = data.get("computers", {})
    bt = c.get("by_type") or {}
    hd = c.get("hard_disk") or {}
    ba = c.get("battery") or {}
    age = c.get("age_distribution") or {}
    make_table(doc, ["No.", "รายละเอียด", "จำนวน", "หมายเหตุ"], [
        ["1", "จำนวนคอมพิวเตอร์ตามสัญญา", c.get("contract_count"), ""],
        ["2", "จำนวนคอมพิวเตอร์ที่ได้รับการบริการ", c.get("serviced_count"), ""],
        ["3", "Desktop", bt.get("desktop"), ""],
        ["4", "Laptop", bt.get("laptop"), ""],
        ["5", "All in One", bt.get("all_in_one"), ""],
        ["6", "MacBook", bt.get("macbook"), ""],
        ["7", "สถานะของ Hard disk — Normal", hd.get("normal"), ""],
        ["8", "สถานะของ Hard disk — Caution", hd.get("caution"),
         "🟡 ตรวจสอบด่วน" if (hd.get("caution") or 0) > 0 else ""],
        ["9", "สถานะของ Battery Laptop — ใช้งานได้ตามปกติ", ba.get("ok"), ""],
        ["10", "สถานะของ Battery Laptop — เสื่อมสภาพ (ไม่เก็บไฟ)", ba.get("degraded"),
         "🟡 แนะนำเปลี่ยน" if (ba.get("degraded") or 0) > 0 else ""],
        ["11", "อายุคอมพิวเตอร์ — น้อยกว่า 1 ปี", age.get("under_1y"), ""],
        ["12", "อายุคอมพิวเตอร์ — 1–2 ปี", age.get("1_2y"), ""],
        ["13", "อายุคอมพิวเตอร์ — 2–4 ปี", age.get("2_4y"), ""],
        ["14", "อายุคอมพิวเตอร์ — 5–7 ปี", age.get("5_7y"), ""],
        ["15", "อายุคอมพิวเตอร์ — มากกว่า 7 ปี", age.get("over_7y"),
         "🔴 วางแผนเปลี่ยน" if (age.get("over_7y") or 0) > 0 else ""],
    ], widths=[Inches(0.5), Inches(3.6), Inches(0.8), Inches(1.37)],
        center_cols=(0, 2))
    if c.get("replacement_recommendation"):
        p = para(doc)
        text(p, "Recommendation: ", size=14, color=dark, bold=True)
        text(p, c["replacement_recommendation"], size=14)
    doc.add_page_break()

    # ---- 3. Software ----
    banner(doc, "Monitoring Report", "Software", bookmark="sec_software", bid=103)
    sw = data.get("software", {})

    def license_rows(items):
        return [[str(i + 1), it.get("name"), fmt_date(it.get("expire")),
                 it.get("amount"), license_status(it.get("expire"))]
                for i, it in enumerate(items) if it.get("name")]

    sub_banner(doc, "Software License — Operating System (OS)")
    make_table(doc, ["No.", "Software Name", "Expire Date", "Amount", "Status"],
               license_rows(sw.get("os_licenses") or []),
               widths=[Inches(0.5), Inches(2.3), Inches(1.05), Inches(0.75),
                       Inches(1.67)], center_cols=(0, 2, 3, 4))
    sub_banner(doc, "Software License — Microsoft Office")
    make_table(doc, ["No.", "Software Name", "Expire Date", "Amount", "Status"],
               license_rows(sw.get("office_licenses") or []),
               widths=[Inches(0.5), Inches(2.3), Inches(1.05), Inches(0.75),
                       Inches(1.67)], center_cols=(0, 2, 3, 4))
    if sw.get("license_alerts"):
        p = para(doc)
        text(p, "⚠️ License alerts: ", size=14, color=red, bold=True)
        text(p, " · ".join(sw["license_alerts"]), size=14, color=red)
    doc.add_page_break()

    # ---- 4. Server ----
    banner(doc, "Monitoring Report", "Server", bookmark="sec_server", bid=104)
    sv = data.get("server", {})
    sub_banner(doc, "Maintenance Server")
    make_table(doc, ["No.", "รายละเอียด", "จำนวน"], [
        ["1", "จำนวน Server ตามสัญญา", sv.get("contract_count")],
        ["2", "จำนวน Server ที่ได้รับการบริการ", sv.get("serviced_count")],
    ], widths=[Inches(0.5), Inches(4.97), Inches(0.8)], center_cols=(0, 2))

    servers = [s for s in sv.get("physical_servers") or [] if s.get("model")]
    if servers:
        sub_banner(doc, "Physical Server Information")
        make_table(doc,
                   ["No.", "Model", "CPU", "RAM (GB)", "Disk", "OS",
                    "Type Server", "Status"],
                   [[str(i + 1), s.get("model"), s.get("cpu"), s.get("ram_gb"),
                     s.get("disk"), s.get("os"), s.get("role"),
                     status_icon(s.get("status"))] for i, s in enumerate(servers)],
                   widths=[Inches(0.35), Inches(0.95), Inches(1.15),
                           Inches(0.55), Inches(1.25), Inches(0.55),
                           Inches(0.85), Inches(0.62)], center_cols=(0, 3))

    sub_banner(doc, "Guest VM Information")
    vms = sv.get("guest_vms") or []
    if vms:
        make_table(doc, ["No.", "Name", "vCPU", "vRAM", "Disk", "OS"],
                   [[str(i + 1), v.get("name"), v.get("vcpu"), v.get("vram"),
                     v.get("disk"), v.get("os")] for i, v in enumerate(vms)],
                   widths=[Inches(0.5), Inches(1.6), Inches(0.8), Inches(0.8),
                           Inches(1.2), Inches(1.37)], center_cols=(0, 2, 3))
    else:
        p = para(doc)
        text(p, "ไม่มี Guest VM ภายใต้สัญญา", size=14, color=gray)

    backup = sv.get("backup", {})
    sub_banner(doc, "Backup")
    kv_table(doc, [
        ("Scheduled backup", status_icon(backup.get("scheduled_status"))),
        ("Last restore test verified",
         backup.get("last_restore_test") or "⚠️ ยังไม่เคยทดสอบ restore"),
    ], bordered=True)
    doc.add_page_break()

    # ---- 5. Firewall / Gateway ----
    banner(doc, "Monitoring Report", "Firewall, Gateway", bookmark="sec_firewall", bid=105)
    sub_banner(doc, "Maintenance Gateway")
    fg = data.get("firewall_gateway", {})
    devices = [d for d in fg.get("devices") or [] if d.get("name")]
    make_table(doc, ["No.", "รายละเอียด", "Firmware", "License Expiry", "สถานะ"],
               [[str(i + 1), d.get("name"), d.get("firmware") or "⚠️ MISSING",
                 d.get("license_expiry") or "⚠️ MISSING",
                 status_icon(d.get("status"))] for i, d in enumerate(devices)],
               widths=[Inches(0.4), Inches(1.6), Inches(1.15), Inches(1.15),
                       Inches(1.97)], center_cols=(0, 2, 3, 4))
    if fg.get("warning_critical_log"):
        p = para(doc)
        text(p, "Warning and Critical Log: ", size=14, color=dark, bold=True)
        text(p, fg["warning_critical_log"], size=14)
    if not any(d.get("status") for d in devices):
        p = para(doc)
        text(p, "🔴 ไม่มีข้อมูลสถานะอุปกรณ์บางรายการ — ควรตรวจสอบและกรอกให้ครบ "
                "ไม่เช่นนั้นลูกค้าจะตั้งคำถามว่ามีการดูแลจริงหรือไม่",
             size=14, color=red, bold=True)
    doc.add_page_break()

    # ---- 6. Network ----
    banner(doc, "Monitoring Report", "Network", bookmark="sec_network", bid=106)
    sub_banner(doc, "Maintenance Network")
    ndevices = [d for d in data.get("network", {}).get("devices") or []
                if d.get("type")]
    make_table(doc, ["No.", "รายละเอียด", "จำนวน", "สถานะ"],
               [[str(i + 1), d.get("type"), d.get("count"),
                 status_icon(d.get("status"))] for i, d in enumerate(ndevices)],
               widths=[Inches(0.5), Inches(2.97), Inches(1.0), Inches(1.8)],
               center_cols=(0, 2, 3))
    nw_log = data.get("network", {}).get("warning_critical_log")
    if nw_log:
        p = para(doc)
        text(p, "Warning and Critical Log: ", size=14, color=dark, bold=True)
        text(p, nw_log, size=14)
    doc.add_page_break()

    # ---- 7. Tickets ----
    banner(doc, "Ticket Report", "Summary", bookmark="sec_ticket", bid=107)
    tk = data.get("tickets", {})
    ts = tk.get("status") or {}
    make_table(doc, ["Type", "Count", "Status", "Count"], [
        ["Incident", tk.get("incident_count", 0), "In progress", ts.get("in_progress", 0)],
        ["Service Request", tk.get("service_request_count", 0), "Pending", ts.get("pending", 0)],
        ["Total", (tk.get("incident_count") or 0) + (tk.get("service_request_count") or 0),
         "Done", ts.get("done", 0)],
    ], widths=[Inches(1.8), Inches(1.3), Inches(1.8), Inches(1.37)],
        center_cols=(1, 3))
    tlist = [t for t in tk.get("list") or [] if t.get("name")]
    if tlist:
        make_table(doc,
                   ["No.", "Name", "Detail", "Status", "Responsible",
                    "Start", "End", "Resolution"],
                   [[str(i + 1), t.get("name"), t.get("detail"), t.get("status"),
                     t.get("responsible"), t.get("start"), t.get("end"),
                     t.get("resolution")] for i, t in enumerate(tlist)],
                   widths=[Inches(0.35), Inches(0.9), Inches(0.8), Inches(0.6),
                           Inches(0.95), Inches(0.75), Inches(0.75),
                           Inches(1.17)], size=12, center_cols=(0, 3, 5, 6))
    else:
        p = para(doc)
        text(p, f"ไม่มี ticket ในเดือน {report.get('month')} "
                f"{report.get('year')} (ตั๋วเดือนก่อนๆ บันทึกแยกที่ history.yaml)",
             size=14, color=gray)
    doc.add_page_break()

    # ---- 8. Scope of Work ----
    banner(doc, "Monthly Report", "Scope of Work Check", bookmark="sec_scope", bid=108)
    sow = [s for s in data.get("scope_of_work") or [] if s.get("item")]
    make_table(doc, ["No.", "Item", "Delivered"],
               [[str(i + 1), s.get("item"),
                 "✅" if s.get("delivered") else "❌"] for i, s in enumerate(sow)],
               widths=[Inches(0.5), Inches(4.67), Inches(1.1)],
               center_cols=(0, 2))
    doc.add_page_break()

    # ---- 9. Recommendations + Sign-off ----
    banner(doc, "Monthly Report", "Recommendations", bookmark="sec_recs", bid=109)
    for r in data.get("recommendations") or []:
        if not r.get("text"):
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        pr = str(r.get("priority", "")).upper()
        pcolor = {"HIGH": red, "MEDIUM": dark, "LOW": gray}.get(pr, ink)
        text(p, f"[{pr}] ", font=en_font, size=14, color=pcolor, bold=True)
        text(p, r["text"], size=14)

    sub_banner(doc, "Sign-off")
    p = para(doc)
    text(p, "Reviewed and acknowledged by "
            f"({report.get('client_name') or '⚠️ MISSING client_name'}):",
         size=14)
    so = data.get("sign_off", {})
    for label, v in (("Name", so.get("client_signer_name")),
                     ("Date", so.get("client_signed_date"))):
        p = para(doc, before=8)
        text(p, f"{label}: ", font=en_font, size=14, color=dark, bold=True)
        text(p, v or "_" * 25, font=en_font, size=14)
    p = para(doc, before=8)
    text(p, "Signature: ", font=en_font, size=14, color=dark, bold=True)
    text(p, "_" * 25, font=en_font, size=14)

    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("data_file", type=Path)
    parser.add_argument("--docx", action="store_true", help="also render a .docx alongside the .md")
    args = parser.parse_args()

    if not args.data_file.exists():
        sys.exit(f"data file not found: {args.data_file}")

    data = load_data(args.data_file)
    md = render_markdown(data)

    out_md = args.data_file.with_suffix(".md")
    out_md.write_text(md, encoding="utf-8")
    print(f"wrote {out_md}")

    if args.docx:
        out_docx = args.data_file.with_suffix(".docx")
        render_docx(data, out_docx)
        print(f"wrote {out_docx}")


if __name__ == "__main__":
    main()
